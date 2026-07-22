from docx import Document as DocxDocument

from src.loaders.base_loader import BaseLoader
from src.core.document import Document


class DOCXLoader(BaseLoader):

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self):
        docx = DocxDocument(self.file_path)
        documents = []

        text = "\n".join(
            paragraph.text
            for paragraph in docx.paragraphs
            if paragraph.text.strip()
        )

        doc = Document(
            content=text,
            metadata={
                "source": self.file_path,
                "page": 1
            }
        )

        documents.append(doc)

        return documents